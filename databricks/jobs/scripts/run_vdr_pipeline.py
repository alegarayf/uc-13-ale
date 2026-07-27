"""
run_vdr_pipeline.py — VDR-integrated diligence pipeline wrapper.

Reads a record from a companies_vdr_history Delta table, runs the full
diligence pipeline (Phase 1-5), and writes results back to the table.

The VDR UI creates records in the table with processing_status='submitted'.
This job picks up a record by id, transitions it through 'processing' → 'done'
(or 'error'), and copies output files to the VDR volume at the path format
required by the frontend:

    /Volumes/rallyday_partners_llc/default/vdr/{company_snake}/{ISO_timestamp}/
        executive_summary.docx
        full_report.docx

Invocation (spark_python_task positional argv):
    run_vdr_pipeline.py <tableName> <id>

Also callable as a Python function::

    from run_vdr_pipeline import run_vdr_pipeline
    result = run_vdr_pipeline("rallyday_partners_llc.default.companies_vdr_history", 42)
"""

import inspect
import json
import os
import re
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path


# ---------------------------------------------------------------------------
# Path helpers (same pattern as run_full_pipeline.py)
# ---------------------------------------------------------------------------

def _find_scripts_dir() -> str:
    here = Path(inspect.getfile(_find_scripts_dir)).resolve().parent
    if (here / "run_full_pipeline.py").exists():
        return str(here)
    for start in (Path.cwd(), here):
        for c in [start, *start.parents]:
            s = c / "jobs" / "scripts"
            if (s / "run_full_pipeline.py").exists():
                return str(s)
    raise RuntimeError(
        "Cannot locate jobs/scripts directory (run_full_pipeline.py not found)."
    )


def _find_repo_root(scripts_dir: str) -> str:
    for c in [Path(scripts_dir), *Path(scripts_dir).parents]:
        if (c / "agents").exists():
            return str(c)
    raise RuntimeError(
        f"Cannot locate repo root (no 'agents' directory found above {scripts_dir!r})."
    )


# ---------------------------------------------------------------------------
# Spark + Delta table helpers
# ---------------------------------------------------------------------------

def _get_spark():
    from pyspark.sql import SparkSession
    spark = SparkSession.getActiveSession()
    if spark is None:
        spark = SparkSession.builder.getOrCreate()
    return spark


def _read_vdr_record(spark, table_name: str, record_id: int) -> dict:
    """Read one row from companies_vdr_history by id. Raises if not found."""
    rows = spark.sql(
        f"SELECT * FROM {table_name} WHERE id = :id",
        args={"id": record_id},
    ).collect()
    if not rows:
        raise ValueError(f"No record found in {table_name} with id={record_id}")
    return rows[0].asDict()


def _update_vdr_record(spark, table_name: str, record_id: int, updates: dict):
    """Update specific columns on the VDR record using parameterized SQL."""
    set_clauses = []
    args = {"id": record_id}
    for i, (col, val) in enumerate(updates.items()):
        param = f"p{i}"
        set_clauses.append(f"{col} = :{param}")
        args[param] = val
    set_sql = ", ".join(set_clauses)
    spark.sql(
        f"UPDATE {table_name} SET {set_sql} WHERE id = :id",
        args=args,
    )


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _to_snake(name: str) -> str:
    return name.strip().lower().replace(" ", "_").replace("/", "_").replace("-", "_")


def _build_output_dir(company_name: str) -> str:
    """Build timestamped output directory path for VDR volume.

    Returns path like /Volumes/rallyday_partners_llc/default/vdr/{snake}/{20260706T134705Z}/
    """
    snake = _to_snake(company_name)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    return f"/Volumes/rallyday_partners_llc/default/vdr/{snake}/{ts}"


# ---------------------------------------------------------------------------
# PDF generation helpers
# ---------------------------------------------------------------------------

def _fmt_dollars(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
        if abs(f) >= 1_000_000:
            return f"${f / 1_000_000:,.1f}M"
        if abs(f) >= 1_000:
            return f"${f / 1_000:,.0f}K"
        return f"${f:,.0f}"
    except Exception:
        return str(v)


def _fmt_pct(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.1f}%"
    except Exception:
        return str(v)


def _md_escape_rl(text: str) -> str:
    """Escape characters that conflict with ReportLab's XML parser."""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    return text


def _md_inline_to_rl(text: str) -> str:
    """Convert basic markdown inline formatting to ReportLab XML tags.

    Order matters: bold before italic, code last so underscores inside
    backtick spans are already inside <font> and won't be re-matched.
    Word-boundary look-arounds on _ prevent matching underscores that are
    part of identifiers (e.g. customer_quality_report).
    """
    text = _md_escape_rl(text)
    # bold (**text**) — must come before single-* italic
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text, flags=re.DOTALL)
    # italic with * — only when not inside a word (**already consumed above**)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text, flags=re.DOTALL)
    # italic with _ — only at non-word boundaries so identifiers are left alone
    text = re.sub(r"(?<!\w)_(?!\s)(.+?)(?<!\s)_(?!\w)", r"<i>\1</i>", text, flags=re.DOTALL)
    # inline code — after bold/italic so backtick content is untouched by above
    text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
    return text


def _rl_append(story: list, rl_text: str, style, raw_line: str) -> None:
    """Append a ReportLab Paragraph, falling back to plain escaped text on parse error."""
    from reportlab.platypus import Paragraph as _Para
    try:
        story.append(_Para(rl_text, style))
    except Exception:
        # If inline formatting produced malformed XML, strip all tags and render plain
        plain = _md_escape_rl(raw_line)
        try:
            story.append(_Para(plain, style))
        except Exception:
            pass  # skip unrenderable line rather than crashing the whole PDF


def _build_full_report_pdf(output_path: str, md_path: str) -> None:
    """Render the markdown diligence memo as a styled PDF using ReportLab."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    NAVY = colors.HexColor("#002855")
    BLUE = colors.HexColor("#0057A8")
    GREY = colors.HexColor("#555555")

    base = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=base["Normal"], fontName="Helvetica-Bold",
                         fontSize=17, textColor=NAVY, spaceBefore=18, spaceAfter=8, leading=22)
    H2 = ParagraphStyle("H2", parent=base["Normal"], fontName="Helvetica-Bold",
                         fontSize=13, textColor=BLUE, spaceBefore=14, spaceAfter=5, leading=17)
    H3 = ParagraphStyle("H3", parent=base["Normal"], fontName="Helvetica-Bold",
                         fontSize=11, textColor=GREY, spaceBefore=10, spaceAfter=3, leading=14)
    BODY = ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica",
                           fontSize=9.5, spaceAfter=4, leading=14)
    BULLET = ParagraphStyle("Bullet", parent=base["Normal"], fontName="Helvetica",
                              fontSize=9.5, leftIndent=16, spaceAfter=2, leading=13,
                              bulletIndent=4)
    MONO = ParagraphStyle("Mono", parent=base["Normal"], fontName="Courier",
                           fontSize=8, spaceAfter=1, leading=11, textColor=GREY)
    BLOCKQUOTE = ParagraphStyle("BQ", parent=base["Normal"], fontName="Helvetica-Oblique",
                                  fontSize=9.5, leftIndent=20, spaceAfter=4, leading=14,
                                  textColor=GREY)

    doc = SimpleDocTemplate(
        output_path, pagesize=letter,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=0.9 * inch,
    )

    story = []
    in_table = False

    for raw in lines:
        stripped = raw.rstrip()

        # Blank line
        if not stripped:
            if in_table:
                in_table = False
            story.append(Spacer(1, 4))
            continue

        # Markdown table rows
        if stripped.startswith("|"):
            in_table = True
            # Skip separator rows (|---|---|)
            if re.match(r"^\|[-| :]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            row_text = " | ".join(_md_escape_rl(c) for c in cells)
            story.append(Paragraph(row_text, MONO))
            continue

        in_table = False

        if stripped.startswith("# "):
            story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=4))
            story.append(Paragraph(_md_inline_to_rl(stripped[2:]), H1))
        elif stripped.startswith("## "):
            story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=3))
            story.append(Paragraph(_md_inline_to_rl(stripped[3:]), H2))
        elif stripped.startswith("### "):
            story.append(Paragraph(_md_inline_to_rl(stripped[4:]), H3))
        elif stripped.startswith(("- ", "* ")):
            _rl_append(story, "• " + _md_inline_to_rl(stripped[2:]), BULLET, stripped)
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            _rl_append(story, "• " + _md_inline_to_rl(text), BULLET, stripped)
        elif stripped.startswith("> "):
            _rl_append(story, _md_inline_to_rl(stripped[2:]), BLOCKQUOTE, stripped)
        elif stripped.startswith("---") or stripped.startswith("==="):
            story.append(HRFlowable(width="100%", thickness=0.5, color=GREY))
        else:
            _rl_append(story, _md_inline_to_rl(stripped), BODY, stripped)

    doc.build(story)


def _build_executive_summary_docx(
    output_path: str,
    company_name: str,
    spark,
    catalog: str = "uc13",
) -> None:
    """Generate a 1-page executive summary Word doc with key financials from Delta tables."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x00, 0x28, 0x55)
    BLUE = RGBColor(0x00, 0x57, 0xA8)
    GREY = RGBColor(0x55, 0x55, 0x55)

    # ── Fetch executive summary text ────────────────────────────────────────
    exec_summary = ""
    try:
        rows = spark.sql(
            f"SELECT executive_summary FROM {catalog}.analysis.diligence_report "
            "WHERE company_name = :c ORDER BY created_at DESC LIMIT 1",
            args={"c": company_name},
        ).collect()
        if rows:
            exec_summary = (rows[0]["executive_summary"] or "").strip()
    except Exception as e:
        exec_summary = f"[Executive summary unavailable: {e}]"

    # ── Fetch financial KPIs from financial_trends ──────────────────────────
    revenue_ttm = "—"
    yoy_growth = "—"
    ebitda_adj = "—"
    try:
        rows = spark.sql(
            f"SELECT revenue_trend_json, ebitda_json FROM {catalog}.analysis.financial_trends "
            "WHERE company_name = :c ORDER BY created_at DESC LIMIT 1",
            args={"c": company_name},
        ).collect()
        if rows:
            rev_trend = json.loads(rows[0]["revenue_trend_json"] or "[]")
            ebitda_data = json.loads(rows[0]["ebitda_json"] or "[]")

            def _period_rank(r):
                p = (r.get("period") or "").upper()
                return 9999 if ("TTM" in p or "LTM" in p) else 0

            sorted_rev = sorted(rev_trend, key=_period_rank, reverse=True)
            if sorted_rev:
                best = sorted_rev[0]
                revenue_ttm = _fmt_dollars(best.get("revenue_stated"))
                yoy_growth = _fmt_pct(best.get("yoy_growth_pct"))

            def _ebitda_rank(r):
                p = (r.get("period") or "").upper()
                v = (r.get("version") or "").lower()
                period_score = 9999 if ("TTM" in p or "LTM" in p) else 0
                version_score = (
                    30 if "pf_adjusted" in v or "clinic_level" in v
                    else 20 if "adjusted" in v
                    else 10 if "reported" in v
                    else 0
                )
                return period_score + version_score

            sorted_ebitda = sorted(ebitda_data, key=_ebitda_rank, reverse=True)
            if sorted_ebitda:
                ebitda_adj = _fmt_dollars(sorted_ebitda[0].get("ebitda_dollars"))
    except Exception as e:
        ebitda_adj = f"[unavailable: {e}]"

    # ── Build Word doc ───────────────────────────────────────────────────────
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(f"Executive Summary — {company_name}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"VDR Diligence  ·  {generated}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = GREY

    doc.add_paragraph()

    # Business Overview heading
    hdr_p = doc.add_paragraph()
    hdr_run = hdr_p.add_run("Business Overview")
    hdr_run.bold = True
    hdr_run.font.size = Pt(12)
    hdr_run.font.color.rgb = BLUE

    # Executive summary blurb
    if exec_summary:
        sentences = re.split(r"(?<=[.!?])\s+", exec_summary)
        blurb = " ".join(sentences[:4]) if len(sentences) >= 4 else exec_summary
        doc.add_paragraph(blurb)
    else:
        doc.add_paragraph("Executive summary not yet generated.")

    doc.add_paragraph()

    # KPIs heading
    kpi_hdr = doc.add_paragraph()
    kpi_run = kpi_hdr.add_run("Key Financial Metrics (TTM)")
    kpi_run.bold = True
    kpi_run.font.size = Pt(12)
    kpi_run.font.color.rgb = BLUE

    # KPI table
    kpi_rows = [
        ("Metric", "Value"),
        ("Revenue TTM", revenue_ttm),
        ("Revenue YoY Growth", yoy_growth),
        ("EBITDA Adjusted", ebitda_adj),
    ]
    tbl = doc.add_table(rows=len(kpi_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(kpi_rows):
        row = tbl.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_run = label_cell.paragraphs[0].add_run(label)
        value_cell.paragraphs[0].add_run(value)
        if i == 0:
            label_run.bold = True
            label_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            value_cell.paragraphs[0].runs[0].bold = True
            value_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for cell in (label_cell, value_cell):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "002855")
                tcPr.append(shd)

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_run = note_p.add_run("Full diligence report: see full_report.docx in this folder.")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = GREY

    # UC Volumes (FUSE) don't support the random-access writes that python-docx's
    # ZipFile backend requires. Write to local disk first, then copy to the volume.
    import tempfile
    import shutil as _shutil
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        _shutil.copy2(tmp_path, output_path)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Main pipeline wrapper
# ---------------------------------------------------------------------------

def run_vdr_pipeline(table_name: str, record_id: int) -> dict:
    """Run the full diligence pipeline for a VDR record.

    Returns::

        {
            "status": "success",
            "company_name": str,
            "output_dir": str,
            "files": [str, ...],
            "pipeline_result": {full pipeline result dict},
        }
    """
    spark = _get_spark()

    # ── Step 1: Read the VDR record ─────────────────────────────────────
    record = _read_vdr_record(spark, table_name, record_id)
    company_name = record["company_name"]
    source_data_location = record.get("source_data_location", "")

    if not company_name:
        raise ValueError(f"Record {record_id} has no company_name")

    print(f"=== VDR Pipeline: {company_name} (id={record_id}) ===")
    print(f"    Source: {source_data_location}")

    # ── Step 2: Update status → processing ──────────────────────────────
    _update_vdr_record(spark, table_name, record_id, {
        "processing_status": "processing",
        "updated_at":        _now_iso(),
        "last_updated_by":   "vdr-backend-ai",
    })

    try:
        # ── Step 3: Set up sys.path and env vars ────────────────────────
        scripts_dir = _find_scripts_dir()
        repo_root = _find_repo_root(scripts_dir)
        for p in (repo_root, scripts_dir):
            if p not in sys.path:
                sys.path.insert(0, p)

        os.environ["sp_company_name"] = company_name

        # ── Step 4: Run the existing pipeline ───────────────────────────
        from agents.shared.agent_base import reset_token_counter, get_token_totals, print_token_summary
        from run_full_pipeline import run_full_pipeline

        reset_token_counter()
        result = run_full_pipeline(company_name=company_name)
        token_totals = get_token_totals()
        print_token_summary()

        # ── Step 5: Check for pipeline success ──────────────────────────
        summary = result.get("summary", {})
        dil_summary = summary.get("diligence", {})
        if dil_summary.get("SUCCESS", 0) == 0 and dil_summary.get("FAILED", 0) > 0:
            raise RuntimeError(
                f"Pipeline completed but all diligence agents failed. "
                f"Summary: {summary}"
            )

        # ── Step 6: Copy outputs to VDR volume ─────────────────────────
        output_dir = _build_output_dir(company_name)

        spark.sql("CREATE VOLUME IF NOT EXISTS rallyday_partners_llc.default.vdr")
        os.makedirs(output_dir, exist_ok=True)

        report_md = result.get("report_md_path")
        report_docx = result.get("report_docx_path")

        files_copied = []

        # full_report.docx — copy the Word doc produced by the orchestrator
        full_report_dst = os.path.join(output_dir, "full_report.docx")
        if report_docx and os.path.exists(report_docx):
            import shutil as _shutil
            _shutil.copy2(report_docx, full_report_dst)
            files_copied.append(full_report_dst)
            print(f"  Copied full_report.docx from: {report_docx}")
        elif report_md and os.path.exists(report_md):
            print(f"  [WARN] .docx not found; falling back to markdown at {report_md}")

        # executive_summary.docx — real Rainmaker-Rev3 one-pager (mine), built
        # via the exec_summary ↔ DAG bridge (replaces the old 1-page placeholder).
        exec_summary_dst = os.path.join(output_dir, "executive_summary.docx")
        print(f"  Building executive_summary.docx (Rev3 one-pager) ...")
        from agents.exec_summary.pipeline_entry import build_exec_summary

        llm_endpoint = os.environ.get("llm_endpoint", "databricks-claude-sonnet-4-6")
        exec_summary_paths = build_exec_summary(
            company_name=company_name,
            catalog="uc13",
            spark=spark,
            llm_endpoint=llm_endpoint,
        )
        tldr_docx = exec_summary_paths.get("tldr_docx")
        if tldr_docx and os.path.exists(tldr_docx):
            import shutil as _shutil
            _shutil.copy2(tldr_docx, exec_summary_dst)
            files_copied.append(exec_summary_dst)
        else:
            print(f"  [WARN] build_exec_summary did not produce a tldr_docx output")

        print(f"  Output files → {output_dir}")
        for f in files_copied:
            print(f"    · {f}")

        # ── Step 7: Update VDR record → done ───────────────────────────
        model_name = os.environ.get("llm_endpoint", "databricks-claude-sonnet-4-6")
        _update_vdr_record(spark, table_name, record_id, {
            "processing_status":  "done",
            "completion_status":  "success",
            "results_location":   output_dir + "/",
            "model_name":         model_name,
            "completion_tokens":  token_totals.get("completion_tokens", 0),
            "prompt_tokens":      token_totals.get("prompt_tokens", 0),
            "total_tokens":       token_totals.get("total_tokens", 0),
            "updated_at":         _now_iso(),
            "last_updated_by":    "vdr-backend-ai",
        })

        print(f"\n=== VDR Pipeline COMPLETE: {company_name} → {output_dir}/ ===")
        return {
            "status":          "success",
            "company_name":    company_name,
            "output_dir":      output_dir,
            "files":           files_copied,
            "pipeline_result": result,
        }

    except Exception as exc:
        # ── Step 8: Update VDR record → error ───────────────────────────
        error_msg = f"{type(exc).__name__}: {exc}"
        tb = traceback.format_exc(limit=6)
        print(f"\n=== VDR Pipeline FAILED: {company_name} ===")
        print(f"  Error: {error_msg}")
        print(tb)

        try:
            _update_vdr_record(spark, table_name, record_id, {
                "processing_status":  "error",
                "completion_status":  "failure",
                "error_message":      error_msg[:4000],
                "updated_at":         _now_iso(),
                "last_updated_by":    "vdr-backend-ai",
            })
        except Exception as update_exc:
            print(f"  [CRITICAL] Failed to update error status: {update_exc}")

        raise


# ---------------------------------------------------------------------------
# Job entry point
# ---------------------------------------------------------------------------

def main():
    argv = sys.argv[1:]

    def _arg(i, key, default=None):
        if i < len(argv) and argv[i]:
            return argv[i]
        return os.environ.get(key, default)

    table_name = _arg(0, "tableName")
    record_id_str = _arg(1, "id")

    if not table_name:
        raise RuntimeError(
            "tableName is required (argv[0] or env var tableName)."
        )
    if not record_id_str:
        raise RuntimeError(
            "id is required (argv[1] or env var id)."
        )

    record_id = int(record_id_str)

    os.environ["tableName"] = table_name
    os.environ["id"] = record_id_str

    print(f"=== VDR Pipeline job: table={table_name}, id={record_id} ===")

    run_vdr_pipeline(table_name, record_id)


if __name__ == "__main__":
    main()
