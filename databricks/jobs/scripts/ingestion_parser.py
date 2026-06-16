"""
03_ingestion_parser.py — Phase 2b: Parsing, chunking, and embedding.

Reads files marked should_parse=true from uc13.classification.doc_relevance,
ordered by priority_tier DESC so Priority Tier documents are processed first.
Parses each file into semantic chunks, generates BGE Large embeddings, and
saves both to Delta tables.

Chunking improvements (vs. original notebook):
  PDF  — Section header carry-forward (no chunk is header-less), document title
          prefix on every chunk: "[Document: {title}] [Section: {header}]\n{text}"
          HTML table elements are converted to pipe-delimited markdown (not
          stripped) so column structure is preserved for the LLM.
          Pages whose ai_parse_document elements are empty figures are passed
          to a vision LLM endpoint for chart/org-chart data extraction.
  Excel — Document + sheet name prefix on every batch. Financial sheets (P&L,
          Balance Sheet, etc.) detect date-like column headers and add a summary
          line "Time periods covered: {cols}" at the top of each chunk.

Phase 2b outputs:
  - Table uc13.ingestion.chunks      (+ source_type: text | table | vision)
  - Table uc13.ingestion.embeddings  (CDF enabled, workstream ARRAY<STRING>,
                                       priority_tier INT, source_type STRING)

Dependencies:
  - uc13.classification.doc_relevance (written by 02_document_classifier.py)
  - Volume files under /Volumes/{catalog}/ingestion/raw_files/{company_name}/
  - python-docx, openpyxl (pre-installed via requirements.txt / cluster init)
  - MLflow endpoint: databricks-bge-large-en
  - Job parameters: sp_company_name, catalog, schema, vision_endpoint (optional)
  - Optional: pymupdf (pip install pymupdf) for vision-based figure extraction
"""

import base64
import csv
import hashlib
import json
import os
import re
import sys
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Secrets / params helpers
# ---------------------------------------------------------------------------

def _get_dbutils():
    """Return the Databricks dbutils object from any execution context.

    Works whether the code runs directly in a notebook cell or is called from
    an imported module (where dbutils is not a direct global but is reachable
    via the IPython user namespace injected by Databricks).
    """
    try:
        return dbutils  # noqa: F821
    except NameError:
        pass
    try:
        import IPython
        user_ns = IPython.get_ipython().user_ns
        if "dbutils" in user_ns:
            return user_ns["dbutils"]
    except Exception:
        pass
    return None


def _load_dotenv_if_local():
    if _get_dbutils() is None:
        try:
            from dotenv import load_dotenv
            load_dotenv()
        except ImportError:
            pass

_load_dotenv_if_local()


def get_secret(key: str) -> str:
    _dbutils = _get_dbutils()
    if _dbutils is not None:
        try:
            return _dbutils.secrets.get("uc13", key)
        except Exception:
            pass
    value = os.environ.get(key)
    if value is None:
        raise RuntimeError(
            f"Secret '{key}' not found. "
            "On Databricks: add it to the 'uc13' secrets scope. "
            "Locally: add it to your .env file or export it as an env var."
        )
    return value


def get_param(key: str, default: str = None) -> str:
    _dbutils = _get_dbutils()
    if _dbutils is not None:
        try:
            value = _dbutils.widgets.get(key)
            if value:
                return value
        except Exception:
            pass
    value = os.environ.get(key, default)
    if value is None:
        raise RuntimeError(
            f"Parameter '{key}' not found. "
            "On Databricks: add it as a job task parameter. "
            "Locally: add it to your .env file or export it as an env var."
        )
    return value


# ---------------------------------------------------------------------------
# Repo root resolver
# ---------------------------------------------------------------------------

def get_current_path():
    try:
        notebook_path = (
            dbutils.notebook.entry_point  # noqa: F821
            .getDbutils()
            .notebook()
            .getContext()
            .notebookPath()
            .get()
        )
        return Path("/Workspace") / notebook_path.lstrip("/")
    except Exception:
        return Path(os.getcwd())


def find_repo_root(marker="agents"):
    current_path = get_current_path()
    if current_path.is_file():
        current_path = current_path.parent
    for path in [current_path, *current_path.parents]:
        if (path / marker).exists():
            return str(path)
    raise RuntimeError(f"Could not find a parent directory containing '{marker}'")


# ---------------------------------------------------------------------------
# Chunk data model
# ---------------------------------------------------------------------------

@dataclass
class Chunk:
    chunk_id: str
    doc_id: str
    file_name: str
    file_type: str
    relative_path: str
    chunk_index: int
    chunk_text: str
    section_header: Optional[str] = None
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    tab: Optional[str] = None
    # "text" | "table" | "vision" — drives source_type column and retrieval routing.
    source_type: str = "text"
    char_count: int = field(init=False)

    def __post_init__(self):
        self.char_count = len(self.chunk_text)


def make_doc_id(path: str) -> str:
    return hashlib.md5(path.encode()).hexdigest()


# ---------------------------------------------------------------------------
# Chunking constants
# ---------------------------------------------------------------------------

MAX_CHUNK_CHARS            = 7_500  # hard ceiling — BGE Large limit with safety buffer
MIN_CHUNK_CHARS            = 150    # drop chunks shorter than this (content only, after prefix)
MIN_VISION_CHUNK_CHARS     = 50     # vision-extracted chart data is naturally shorter
CHUNK_OVERLAP_CHARS        = 200    # overlap for narrative content (PDF prose, Word)
OPERATIONAL_ROWS_PER_CHUNK = 50    # non-financial Excel/CSV rows per chunk
FINANCIAL_LINES_PER_CHUNK  = 30    # financial sheet line items per chunk


# ---------------------------------------------------------------------------
# Shared regex / helpers
# ---------------------------------------------------------------------------

_SKIP_ELEMENT_TYPES  = {"page_footer", "page_number", "header"}
_HEADER_ELEMENT_TYPES = {"section_header", "title"}

_BAD_TITLE_WORDS = re.compile(
    r"^(disclaimer|confidential|important\s+notice|table\s+of\s+contents"
    r"|cover\s+page|this\s+document|prepared\s+by|strictly\s+private)$",
    re.IGNORECASE,
)

_HTML_TAG_RE = re.compile(r"<[^>]+>")

_DATE_HEADER_RE = re.compile(
    r"\d{4}|\w+ \d{4}|Q\d \d{4}|H[12] \d{4}|TTM|LTM|Budget|Actual|Forecast",
    re.IGNORECASE,
)

_FINANCIAL_SHEET_RE = re.compile(
    r"p&l|profit.loss|income|balance sheet|cash flow|addback|ebitda"
    r"|revenue|forecast|budget|model|cogs|margin",
    re.IGNORECASE,
)

# Broader pattern for PDF section headers.  Matches:
#   - standalone financial data keywords (p&l, ebitda, revenue, etc.)
#   - "financials" as a bare section title (common CIM heading)
#   - "historical financials" / "key financials" compound phrases
#   - "financial <data-word>" compounds (financial performance, overview, etc.)
# Does NOT match single-word "financial" to avoid triggering on legal headers
# like "Financial Representations" or "Financial Covenants" — those sections
# contain dense text pages that won't pass the image-page detection anyway,
# but excluding them keeps false vision calls to a minimum.
_PDF_FIN_SECTION_RE = re.compile(
    r"p&l|profit.loss|income statement|balance sheet|cash flow|addback|ebitda"
    r"|revenue|forecast|budget|cogs|margin"
    r"|\bfinancials\b"
    r"|(?:historical|key|selected|summary)\s+financials?"
    r"|financial\s+(?:performance|overview|results|highlights|summary"
    r"|projections?|history|statements?|trends?|model|data|information|metrics)",
    re.IGNORECASE,
)

_PREFIX_RE = re.compile(r"^\[Document:.*?\](\s*\[.*?\])*\s*", re.DOTALL)


def _strip_html(text: str) -> str:
    text = _HTML_TAG_RE.sub(" ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# HTML table → markdown (preserves column structure for the LLM)
# ---------------------------------------------------------------------------

class _HTMLTableParser(HTMLParser):
    """Minimal HTML parser that extracts rows and cells from a table element."""
    def __init__(self):
        super().__init__()
        self.rows: list = []
        self._row: list = []
        self._cell: list = []
        self._in_cell = False

    def handle_starttag(self, tag, attrs):
        if tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell = []
            self._in_cell = True

    def handle_endtag(self, tag):
        if tag == "tr":
            if self._row:
                self.rows.append(self._row[:])
            self._row = []
        elif tag in ("td", "th"):
            self._row.append(" ".join(self._cell).strip())
            self._cell = []
            self._in_cell = False

    def handle_data(self, data):
        if self._in_cell:
            text = data.strip()
            if text:
                self._cell.append(text)


def _html_table_to_markdown(html: str) -> str:
    """Convert an HTML table to pipe-delimited markdown.

    Preserving the column header → value relationship is critical for tables
    where column headers carry the key meaning (e.g. owner names as columns in
    an ownership table, or 'North America | Global' headcount columns).  The
    plain _strip_html fallback loses this entirely.

    Falls back to _strip_html if the HTML contains no <tr> elements or only
    one row (nothing to structure).  This makes it safe to call on any content
    regardless of whether it is actually a well-formed HTML table.
    """
    parser = _HTMLTableParser()
    try:
        parser.feed(html)
    except Exception:
        return _strip_html(html)

    rows = parser.rows
    if len(rows) < 2:
        return _strip_html(html)

    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 0:
        return _strip_html(html)

    padded = [r + [""] * (max_cols - len(r)) for r in rows]
    lines = ["| " + " | ".join(padded[0]) + " |",
             "|" + "|".join(["---"] * max_cols) + "|"]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _is_financial_sheet(sheet_name: str, file_name: str) -> bool:
    return bool(
        _FINANCIAL_SHEET_RE.search(sheet_name)
        or _FINANCIAL_SHEET_RE.search(file_name)
    )


def _expand_merged_cells(ws) -> None:
    """Copy the top-left cell value to every cell in each merged range, then unmerge.

    Without this, openpyxl returns None for all non-top-left cells of a merge.
    Critical for financial P&L sheets where period headers (2020A, 2021A…) sit
    under spanning headers like "Actual" or "Projected" — those spanning cells
    must be propagated so column-index mapping in _chunk_financial_sheet is correct.

    Must be called on a non-read_only worksheet (read_only=True worksheets do not
    expose the merged_cells attribute).
    """
    try:
        for merge_range in list(ws.merged_cells.ranges):
            top_left_value = ws.cell(merge_range.min_row, merge_range.min_col).value
            ws.unmerge_cells(str(merge_range))
            for row in range(merge_range.min_row, merge_range.max_row + 1):
                for col in range(merge_range.min_col, merge_range.max_col + 1):
                    ws.cell(row, col).value = top_left_value
    except Exception:
        pass  # merged_cells unavailable on read-only sheets — skip silently


def _is_valid_chunk(chunk_text: str) -> bool:
    """Return False for chunks below MIN_CHUNK_CHARS or above MAX_CHUNK_CHARS.

    Oversized chunks are dropped (with a warning) rather than forwarded to
    get_embeddings_batch, which raises ValueError on them.
    """
    stripped = chunk_text.strip()
    content  = _PREFIX_RE.sub("", stripped).strip()
    if len(content) < MIN_CHUNK_CHARS:
        return False
    if len(stripped) > MAX_CHUNK_CHARS:
        print(f"  ⚠ Dropping oversized chunk ({len(stripped):,} chars — exceeds MAX_CHUNK_CHARS {MAX_CHUNK_CHARS:,}).")
        return False
    return True


def _split_long_text(text: str, prefix: str) -> list[str]:
    """Split text that would exceed MAX_CHUNK_CHARS into sub-chunks with overlap.

    Each sub-chunk carries the same prefix so retrieval context is never lost.
    Splits on sentence boundaries first, then paragraph breaks, then words.
    """
    if len(prefix) + len(text) <= MAX_CHUNK_CHARS:
        return [f"{prefix}\n\n{text}"]

    separators   = [". ", "\n\n", "\n", " "]
    effective_max = MAX_CHUNK_CHARS - len(prefix) - 10
    results: list[str] = []

    while len(text) > effective_max:
        split_at = effective_max
        for sep in separators:
            pos = text.rfind(sep, 0, effective_max)
            if pos > effective_max * 0.5:
                split_at = pos + len(sep)
                break
        results.append(f"{prefix}\n\n{text[:split_at].strip()}")
        overlap_start = max(0, split_at - CHUNK_OVERLAP_CHARS)
        text = text[overlap_start:]

    if text.strip():
        results.append(f"{prefix}\n\n{text.strip()}")
    return results


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------


def parse_pdf(
    file_path: str,
    doc_id: str,
    file_name: str,
    spark,
    vision_endpoint: Optional[str] = None,
) -> list[Chunk]:
    """Parse a PDF using Databricks ai_parse_document.

    - Section header carry-forward: no chunk is ever context-free.
    - Table elements are converted to pipe-delimited markdown (source_type='table')
      so column headers remain meaningful to the LLM.
    - Pages that have empty figure elements (charts, diagrams, org charts) are
      optionally passed to a vision LLM if vision_endpoint is provided
      (source_type='vision').  Requires pymupdf.
    - Long sections are split with CHUNK_OVERLAP_CHARS overlap.
    - Every chunk gets: [Document: {title}] [Section: {header}]
    """
    try:
        _df = spark.sql(f"""
            SELECT to_json(ai_parse_document(
                content,
                map('version', '2.0')
            )) AS parsed
            FROM read_files('{file_path}', format => 'binaryFile')
        """)
        _rows = _df.collect()
        if not _rows:
            print(f"  ⚠ Skipped (empty/unreadable file): {file_name}")
            return []
        row = _rows[0]["parsed"]

        result   = json.loads(row)
        elements = result["document"]["elements"]

        # Title: skip boilerplate cover/disclaimer pages.
        doc_title = Path(file_name).stem
        for el in elements:
            if el.get("type") == "title":
                candidate = _strip_html(el.get("content", "")).strip()
                if candidate and not _BAD_TITLE_WORDS.match(candidate):
                    doc_title = candidate
                    break

        chunks: list[Chunk] = []
        current_header: Optional[str] = None
        last_known_header: Optional[str] = None
        current_texts: list[str] = []
        current_pages: list[int] = []
        chunk_index = 0
        # Maps 0-indexed page_id → section header at that point, for vision extraction.
        figure_page_header_map: dict = {}

        def _make_prefix() -> str:
            h = current_header or last_known_header or "Document body"
            return f"[Document: {doc_title}] [Section: {h}]"

        def flush_prose():
            nonlocal chunk_index
            if not current_texts:
                return
            text   = "\n".join(current_texts).strip()
            prefix = _make_prefix()
            for ct in _split_long_text(text, prefix):
                if _is_valid_chunk(ct):
                    h = current_header or last_known_header or "Document body"
                    chunks.append(Chunk(
                        chunk_id=str(uuid.uuid4()),
                        doc_id=doc_id,
                        file_name=file_name,
                        file_type="pdf",
                        relative_path=file_path,
                        chunk_index=chunk_index,
                        chunk_text=ct,
                        section_header=h,
                        page_start=min(current_pages) + 1 if current_pages else None,
                        page_end=max(current_pages) + 1 if current_pages else None,
                        source_type="text",
                    ))
                    chunk_index += 1

        for el in elements:
            el_type     = el.get("type", "")
            raw_content = el.get("content", "")
            content     = _strip_html(raw_content).strip()
            # page_id is nested inside bbox[0] in ai_parse_document v2.0 output,
            # not at the top level of the element.
            _bbox   = el.get("bbox") or []
            page_id = _bbox[0].get("page_id") if _bbox else None

            # Figure elements: content is always empty in v2.0 but ai_parse_document
            # provides a free text description of the figure.  Store the description
            # as the chunk text so retrieval can surface figure context even without
            # the vision LLM.  Also queue the page for vision extraction if enabled.
            if el_type == "figure":
                description = (el.get("description") or "").strip()
                if page_id is not None:
                    figure_page_header_map.setdefault(
                        page_id,
                        current_header or last_known_header or "Document body",
                    )
                if description:
                    h      = current_header or last_known_header or "Document body"
                    prefix = f"[Document: {doc_title}] [Section: {h}] [Figure]"
                    ct     = f"{prefix}\n\n{description}"
                    if len(description) >= MIN_VISION_CHUNK_CHARS:
                        chunks.append(Chunk(
                            chunk_id=str(uuid.uuid4()),
                            doc_id=doc_id,
                            file_name=file_name,
                            file_type="pdf",
                            relative_path=file_path,
                            chunk_index=chunk_index,
                            chunk_text=ct,
                            section_header=h,
                            page_start=page_id + 1 if page_id is not None else None,
                            page_end=page_id + 1 if page_id is not None else None,
                            source_type="vision",
                        ))
                        chunk_index += 1
                continue

            if el_type in _SKIP_ELEMENT_TYPES or not content:
                continue

            if el_type in _HEADER_ELEMENT_TYPES:
                flush_prose()
                current_header    = content
                last_known_header = content
                current_texts     = []
                current_pages     = []

            elif el_type == "table":
                # Convert HTML to markdown so column headers stay meaningful.
                # _strip_html would turn "| Owner A | Owner B |" → "Owner A Owner B"
                # losing the column→value relationship entirely.
                flush_prose()
                current_texts = []
                current_pages = []
                table_md = _html_table_to_markdown(raw_content)
                if not table_md.strip():
                    continue
                prefix = _make_prefix()
                h      = current_header or last_known_header or "Document body"
                for ct in _split_long_text(table_md, prefix):
                    if _is_valid_chunk(ct):
                        chunks.append(Chunk(
                            chunk_id=str(uuid.uuid4()),
                            doc_id=doc_id,
                            file_name=file_name,
                            file_type="pdf",
                            relative_path=file_path,
                            chunk_index=chunk_index,
                            chunk_text=ct,
                            section_header=h,
                            page_start=page_id + 1 if page_id is not None else None,
                            page_end=page_id + 1 if page_id is not None else None,
                            source_type="table",
                        ))
                        chunk_index += 1

            else:
                current_texts.append(content)
                if page_id is not None:
                    current_pages.append(page_id)

        flush_prose()

        # Image-page detection: find pages that ai_parse_document couldn't extract
        # real content from, and queue them for the vision LLM.
        #
        # Key insight: ai_parse_document may still emit page_footer / page_number
        # elements for image-only pages (e.g. a page whose only text is "47").
        # Counting those as "content" makes sparse image pages look populated.
        # Fix: count only meaningful element types (not skip types, not empty
        # figures).  Any page with < 30 meaningful chars is flagged for vision.
        # Pages completely absent from the dict (no elements at all) use -1 as
        # default, which also satisfies < 30.  No financial-section requirement —
        # the vision LLM returns NO_DATA for non-financial pages, so the only
        # cost of a false positive is one extra API call.
        if vision_endpoint:
            _page_meaningful_chars: dict[int, int] = {}
            _page_last_header: dict[int, str] = {}
            _last_any_hdr = last_known_header or "Document"

            for el in elements:
                _bbox    = el.get("bbox") or []
                _pid     = _bbox[0].get("page_id") if _bbox else None
                _el_type = el.get("type", "")
                _content = _strip_html(el.get("content", "")).strip()
                if _el_type in _HEADER_ELEMENT_TYPES:
                    _last_any_hdr = _content
                if _pid is not None:
                    _page_last_header[_pid] = _last_any_hdr
                    _is_skip = _el_type in _SKIP_ELEMENT_TYPES
                    _is_empty_fig = (_el_type == "figure" and not _content)
                    if not _is_skip and not _is_empty_fig:
                        _page_meaningful_chars[_pid] = (
                            _page_meaningful_chars.get(_pid, 0) + len(_content)
                        )
                    elif _pid not in _page_meaningful_chars:
                        # Register the page so gaps between it and real pages
                        # don't widen the scan range unnecessarily.
                        _page_meaningful_chars[_pid] = 0

            if _page_meaningful_chars:
                _min_pid   = min(_page_meaningful_chars.keys())
                _max_pid   = max(_page_meaningful_chars.keys())
                _carry_hdr = last_known_header or "Document"
                for _pid in range(_min_pid, _max_pid + 1):
                    if _pid in _page_last_header:
                        _carry_hdr = _page_last_header[_pid]
                    # -1 default = page absent from elements entirely (also < 30)
                    if _page_meaningful_chars.get(_pid, -1) < 30:
                        figure_page_header_map.setdefault(_pid, _carry_hdr)

        # Vision extraction: render figure-heavy pages and extract chart/org data.
        if vision_endpoint and figure_page_header_map:
            vision_chunks = _extract_figure_pages_with_vision(
                file_path=file_path,
                doc_id=doc_id,
                file_name=file_name,
                doc_title=doc_title,
                figure_page_header_map=figure_page_header_map,
                start_chunk_index=chunk_index,
                vision_endpoint=vision_endpoint,
            )
            chunks.extend(vision_chunks)

        table_count  = sum(1 for c in chunks if c.source_type == "table")
        vision_count = sum(1 for c in chunks if c.source_type == "vision")
        print(
            f"  ✓ {file_name}: {len(chunks)} PDF chunks"
            f" ({table_count} table, {vision_count} vision,"
            f" {len(figure_page_header_map)} figure pages detected)"
        )
        return chunks

    except Exception as exc:
        print(f"  ✗ {file_name}: {exc}")
        return []


# ---------------------------------------------------------------------------
# Vision-based figure extraction
# ---------------------------------------------------------------------------

def _extract_figure_pages_with_vision(
    file_path: str,
    doc_id: str,
    file_name: str,
    doc_title: str,
    figure_page_header_map: dict,
    start_chunk_index: int,
    vision_endpoint: str,
) -> list[Chunk]:
    """Render figure-heavy PDF pages and extract data via a vision LLM endpoint.

    Fires for any page that ai_parse_document returned empty figure elements on —
    covers pie/donut charts (payor mix, referral sources, tenure), bar charts,
    org charts, infographic panels, and any other visual-only content.

    Requires PyMuPDF (pip install pymupdf).  Falls back gracefully to an empty
    list if the library is missing, so ingestion continues without vision data.

    source_type is set to 'vision' on every chunk produced here so downstream
    retrieval can distinguish vision-extracted content from parsed text.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print(
            f"  ⚠ PyMuPDF not installed — vision extraction skipped for {file_name}.\n"
            "    Add pymupdf to the cluster init script or run: %pip install pymupdf"
        )
        return []

    import mlflow.deployments

    _VISION_PROMPT = (
        "This page is from a financial due diligence or business overview document. "
        "Extract ALL data from any charts, graphs, diagrams, org charts, or tables visible.\n"
        "• Pie / donut / bar charts: list each label and its value or percentage, "
        "one per line in the format 'Label: Value'.\n"
        "• Org charts: list each person as 'Name: [full name] | Title: [job title]'.\n"
        "• Data tables: reproduce the header row and all data rows using ' | ' as the "
        "column separator.\n"
        "• Callout boxes or annotations: reproduce the text verbatim.\n"
        "If this page contains no extractable data respond with exactly: NO_DATA\n"
        "Return ONLY the extracted data. No descriptions, no commentary."
    )

    # Financial-specific vision prompt — used when section_header signals a P&L,
    # income statement, EBITDA bridge, or balance sheet page.  Produces column-
    # aligned tabular output optimised for LLM financial extraction.
    _VISION_PROMPT_FINANCIAL = (
        "This page contains a financial statement table (P&L, income statement, "
        "EBITDA bridge, balance sheet, or similar) from a business overview or "
        "due diligence document.\n"
        "Extract ALL data from the table using this exact format:\n"
        "• First output the column header row: Label | 2020A | 2021A | 2022A | ...\n"
        "• Then one line per row item: Revenue | 8,955 | 14,176 | 20,846 | ...\n"
        "• Preserve parentheses for negative values: (342)\n"
        "• Include margin % rows (e.g. 'Margin | 42.1% | 44.3% | ...')\n"
        "• Include growth % rows (e.g. 'Growth | N/A | 58.3% | ...')\n"
        "• If the table has multiple named sections (Revenue, Gross Profit, EBITDA), "
        "separate them with a blank line and the section label.\n"
        "If this page contains no financial table respond with exactly: NO_DATA\n"
        "Return ONLY the table data. No descriptions, no preamble, no commentary."
    )

    chunks: list[Chunk] = []
    chunk_index = start_chunk_index

    try:
        pdf_doc = fitz.open(file_path)
    except Exception as exc:
        print(f"  ⚠ PyMuPDF could not open {file_name}: {exc}")
        return []

    client = mlflow.deployments.get_deploy_client("databricks")

    for page_id, section_header in sorted(figure_page_header_map.items()):
        if page_id >= len(pdf_doc):
            continue
        try:
            page    = pdf_doc[page_id]
            mat     = fitz.Matrix(150 / 72, 150 / 72)  # 150 DPI — quality vs. token cost
            pix     = page.get_pixmap(matrix=mat)
            img_b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")

            # Use the financial-specific prompt when the section header signals a
            # P&L or income statement page — produces column-aligned tabular output
            # that the financial trends agent can extract directly.
            _is_fin_section = bool(_PDF_FIN_SECTION_RE.search(section_header or ""))
            _active_prompt  = _VISION_PROMPT_FINANCIAL if _is_fin_section else _VISION_PROMPT

            response = client.predict(
                endpoint=vision_endpoint,
                inputs={
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{img_b64}"
                                    },
                                },
                                {"type": "text", "text": _active_prompt},
                            ],
                        }
                    ],
                    "max_tokens": 2000,
                },
            )

            text = (
                (response.get("choices") or [{}])[0]
                .get("message", {})
                .get("content", "")
                .strip()
            )

            if not text or text.upper() == "NO_DATA":
                continue

            prefix = (
                f"[Document: {doc_title}] [Section: {section_header}]"
                f" [Page: {page_id + 1}] [Source: vision-extracted figure]"
            )

            for ct in _split_long_text(text, prefix):
                content_only = _PREFIX_RE.sub("", ct).strip()
                if len(content_only) >= MIN_VISION_CHUNK_CHARS:
                    chunks.append(Chunk(
                        chunk_id=str(uuid.uuid4()),
                        doc_id=doc_id,
                        file_name=file_name,
                        file_type="pdf",
                        relative_path=file_path,
                        chunk_index=chunk_index,
                        chunk_text=ct,
                        section_header=section_header,
                        page_start=page_id + 1,
                        page_end=page_id + 1,
                        source_type="vision",
                    ))
                    chunk_index += 1

        except Exception as exc:
            print(f"  ⚠ Vision extraction failed for {file_name} page {page_id + 1}: {exc}")

    pdf_doc.close()

    if chunks:
        print(
            f"  ✓ Vision extracted {len(chunks)} chunk(s) from "
            f"{len(figure_page_header_map)} figure page(s) in {file_name}"
        )
    return chunks


# ---------------------------------------------------------------------------
# Excel parser
# ---------------------------------------------------------------------------

def _chunk_financial_sheet(
    ws, sheet_name: str, file_name: str, doc_id: str,
    file_path: str, start_index: int,
) -> list[Chunk]:
    """Transposed line-item chunking for financial sheets.

    Detects the header row (first row with ≥2 date-like values), then emits
    one chunk per group of FINANCIAL_LINES_PER_CHUNK line items, each prefixed
    with the period labels and formatted as:
        {label}: {col1}=X | {col2}=Y | ...
    """
    all_rows = []
    for row in ws.iter_rows(values_only=True):
        vals = [str(c).strip() if c is not None else "" for c in row]
        if any(v for v in vals):
            all_rows.append(vals)

    if not all_rows:
        return []

    # Detect header row: first row with ≥2 date-like column values.
    header_row_idx = None
    period_headers: list[str] = []
    for i, row in enumerate(all_rows):
        date_matches = [v for v in row[1:] if _DATE_HEADER_RE.search(v)]
        if len(date_matches) >= 2:
            header_row_idx = i
            period_headers = row
            break

    if header_row_idx is None:
        # No date header found — fall back to operational chunking.
        return _chunk_operational_sheet(
            all_rows, sheet_name, file_name, doc_id, file_path, start_index
        )

    data_rows = all_rows[header_row_idx + 1:]
    period_label_line = "Time periods: " + " | ".join(
        h for h in period_headers[1:] if h
    )

    chunks: list[Chunk] = []
    chunk_index = start_index
    line_buffer: list[str] = []

    def flush_financial(section_label: str):
        nonlocal chunk_index
        if not line_buffer:
            return
        prefix = (
            f"[Document: {file_name}] [Sheet: {sheet_name}]"
            f" [Section: {section_label}]"
        )
        body = period_label_line + "\n" + "\n".join(line_buffer)
        for ct in _split_long_text(body, prefix):
            if _is_valid_chunk(ct):
                chunks.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    file_name=file_name,
                    file_type="xlsx",
                    relative_path=file_path,
                    chunk_index=chunk_index,
                    chunk_text=ct,
                    section_header=f"{sheet_name} — {section_label}",
                    tab=sheet_name,
                ))
                chunk_index += 1
        line_buffer.clear()

    current_section = "Summary"
    _SECTION_KEYWORDS = re.compile(
        r"revenue|gross profit|ebitda|net income|operating|headcount"
        r"|total|expense|cost|margin|cash|debt|equity|capex",
        re.IGNORECASE,
    )

    for row in data_rows:
        label  = row[0] if row else ""
        values = row[1:len(period_headers)]

        # Section break: non-empty label, all period values empty/zero.
        all_empty = all(not v or v in ("0", "0.0", "0.00") for v in values)
        if label and all_empty and _SECTION_KEYWORDS.search(label):
            flush_financial(current_section)
            current_section = label
            continue

        # Skip spacer rows.
        if not label and all_empty:
            continue

        period_vals = " | ".join(
            f"{period_headers[i+1]}={values[i]}"
            for i in range(min(len(period_headers) - 1, len(values)))
            if values[i] and values[i] not in ("0", "0.0", "")
        )
        if period_vals:
            line_buffer.append(f"{label}: {period_vals}")

        # Flush on line count OR char budget (whichever comes first).
        buffer_chars = sum(len(l) for l in line_buffer) + len(period_label_line)
        if len(line_buffer) >= FINANCIAL_LINES_PER_CHUNK or buffer_chars >= MAX_CHUNK_CHARS - 300:
            flush_financial(current_section)

    flush_financial(current_section)
    return chunks


def _chunk_operational_sheet(
    all_rows: list[list[str]], sheet_name: str, file_name: str,
    doc_id: str, file_path: str, start_index: int,
) -> list[Chunk]:
    """Row-block chunking for operational sheets (KPI, headcount, customer lists)."""
    if not all_rows:
        return []

    headers   = all_rows[0]
    data_rows = all_rows[1:]
    chunks: list[Chunk] = []
    chunk_index = start_index
    current_section = "Data"
    row_buffer: list[str] = []
    row_start  = 1

    def flush_operational():
        nonlocal chunk_index, row_start
        if not row_buffer:
            return
        end_row  = row_start + len(row_buffer) - 1
        prefix   = (
            f"[Document: {file_name}] [Sheet: {sheet_name}]"
            f" [Section: {current_section} Rows {row_start}–{end_row}]"
        )
        body = "Columns: " + " | ".join(h for h in headers if h) + "\n" + "\n".join(row_buffer)
        for ct in _split_long_text(body, prefix):
            if _is_valid_chunk(ct):
                chunks.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    file_name=file_name,
                    file_type="xlsx",
                    relative_path=file_path,
                    chunk_index=chunk_index,
                    chunk_text=ct,
                    section_header=f"{sheet_name} — {current_section}",
                    tab=sheet_name,
                ))
                chunk_index += 1
        row_buffer.clear()
        row_start = end_row + 1

    for i, row in enumerate(data_rows):
        label  = row[0] if row else ""
        others = row[1:] if len(row) > 1 else []
        all_empty_others = all(not v for v in others)

        # Section label row: non-empty first column, all others empty.
        if label and all_empty_others:
            flush_operational()
            current_section = label
            continue

        row_str = " | ".join(
            f"{headers[j]}: {row[j]}"
            for j in range(min(len(headers), len(row)))
            if row[j].strip()
        )
        if row_str:
            row_buffer.append(row_str)

        if len(row_buffer) >= OPERATIONAL_ROWS_PER_CHUNK:
            flush_operational()

    flush_operational()
    return chunks


def parse_excel(file_path: str, doc_id: str, file_name: str) -> list[Chunk]:
    """Parse an Excel workbook using sheet-type-aware chunking.

    Financial sheets (P&L, Balance Sheet, EBITDA, Forecast, etc.) use transposed
    line-item format — one chunk per group of rows with all period values inline.
    Operational sheets (KPI, headcount, customer lists) use section-aware row batches.
    Every chunk carries [Document][Sheet][Section] prefix and respects MAX_CHUNK_CHARS.

    Loading notes:
      - read_only=False: required to access merged_cells for _expand_merged_cells().
        read_only=True streaming mode does not expose the merged_cells attribute,
        causing merged period headers (e.g. "2020A" spanning two rows) to appear
        as None in all non-top-left cells.
      - data_only=True: returns cached formula results rather than formula strings,
        so cells like =SUM(...) yield their numeric value.
      - Performance: for typical PE financial models (<50 MB) the in-memory load
        is fast; for very large workbooks consider increasing cluster memory.
    """
    import openpyxl

    chunks: list[Chunk] = []

    try:
        wb = openpyxl.load_workbook(file_path, read_only=False, data_only=True)
        chunk_index = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Expand merged cells before row iteration so period headers and
            # section labels propagate correctly across all spanned columns.
            _expand_merged_cells(ws)

            if _is_financial_sheet(sheet_name, file_name):
                new_chunks = _chunk_financial_sheet(
                    ws, sheet_name, file_name, doc_id, file_path, chunk_index
                )
            else:
                all_rows = []
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c).strip() if c is not None else "" for c in row]
                    if any(v for v in vals):
                        all_rows.append(vals)
                new_chunks = _chunk_operational_sheet(
                    all_rows, sheet_name, file_name, doc_id, file_path, chunk_index
                )

            chunks.extend(new_chunks)
            chunk_index += len(new_chunks)

        wb.close()
        print(f"  ✓ {file_name}: {len(chunks)} Excel chunks")
    except Exception as exc:
        print(f"  ✗ {file_name}: {exc}")

    return chunks


# ---------------------------------------------------------------------------
# Word parser
# ---------------------------------------------------------------------------

def parse_word(file_path: str, doc_id: str, file_name: str) -> list[Chunk]:
    """Parse a Word document, splitting on Heading styles.

    - First Heading 1 (or filename stem) used as document title.
    - Every chunk: [Document: {title}] [Section: {header}]
    - Long sections split with CHUNK_OVERLAP_CHARS overlap.
    """
    from docx import Document

    chunks: list[Chunk] = []
    chunk_index = 0
    current_header: Optional[str] = None
    current_texts: list[str] = []

    def flush_word():
        nonlocal chunk_index
        if not current_texts:
            return
        text   = "\n".join(current_texts).strip()
        header = current_header or "Document body"
        prefix = f"[Document: {doc_title}] [Section: {header}]"
        for ct in _split_long_text(text, prefix):
            if _is_valid_chunk(ct):
                chunks.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    file_name=file_name,
                    file_type="docx",
                    relative_path=file_path,
                    chunk_index=chunk_index,
                    chunk_text=ct,
                    section_header=header,
                ))
                chunk_index += 1

    try:
        doc = Document(file_path)

        # Extract document title from the first Heading 1, fall back to filename.
        doc_title = Path(file_name).stem
        for para in doc.paragraphs:
            if para.style.name == "Heading 1" and para.text.strip():
                doc_title = para.text.strip()
                break

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                flush_word()
                current_header = text
                current_texts  = []
            else:
                current_texts.append(text)
        flush_word()
        print(f"  ✓ {file_name}: {len(chunks)} Word chunks")
    except Exception as exc:
        print(f"  ✗ {file_name}: {exc}")

    return chunks


# ---------------------------------------------------------------------------
# CSV parser
# ---------------------------------------------------------------------------

def parse_csv(file_path: str, doc_id: str, file_name: str) -> list[Chunk]:
    """Parse a CSV file.

    Financial CSVs (detected by filename keywords + date-like column headers)
    use the transposed line-item format. All others use OPERATIONAL_ROWS_PER_CHUNK
    row batches. Every chunk carries [Document][Section] prefix.
    """
    chunks: list[Chunk] = []
    chunk_index = 0

    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
            rows = list(csv.reader(fh))
        if not rows:
            return chunks

        headers = rows[0]
        data    = rows[1:]

        # Check if this is a financial CSV.
        date_cols = [h for h in headers if _DATE_HEADER_RE.search(h)]
        if _is_financial_sheet(file_name, "") and len(date_cols) >= 2:
            # Transposed line-item format — reuse operational helper after
            # building all_rows as a list-of-lists.
            all_rows_typed = [headers] + [[str(v) for v in r] for r in data]
            # Delegate to a stub that mimics the Excel financial chunker.
            period_label_line = "Time periods: " + " | ".join(date_cols)
            line_buffer: list[str] = []

            def flush_csv_financial(section: str):
                nonlocal chunk_index
                if not line_buffer:
                    return
                prefix = f"[Document: {file_name}] [Section: {section}]"
                body   = period_label_line + "\n" + "\n".join(line_buffer)
                for ct in _split_long_text(body, prefix):
                    if _is_valid_chunk(ct):
                        chunks.append(Chunk(
                            chunk_id=str(uuid.uuid4()),
                            doc_id=doc_id,
                            file_name=file_name,
                            file_type="csv",
                            relative_path=file_path,
                            chunk_index=chunk_index,
                            chunk_text=ct,
                            section_header=section,
                        ))
                        chunk_index += 1
                line_buffer.clear()

            current_section = "Summary"
            for row in data:
                label  = row[0] if row else ""
                values = row[1:len(headers)]
                all_empty = all(not v or v in ("0", "0.0", "") for v in values)
                if label and all_empty:
                    flush_csv_financial(current_section)
                    current_section = label
                    continue
                period_vals = " | ".join(
                    f"{headers[i+1]}={values[i]}"
                    for i in range(min(len(headers) - 1, len(values)))
                    if values[i] and values[i] not in ("0", "0.0", "")
                )
                if period_vals:
                    line_buffer.append(f"{label}: {period_vals}")
                buffer_chars = sum(len(l) for l in line_buffer) + len(period_label_line)
                if len(line_buffer) >= FINANCIAL_LINES_PER_CHUNK or buffer_chars >= MAX_CHUNK_CHARS - 300:
                    flush_csv_financial(current_section)
            flush_csv_financial(current_section)

        else:
            # Operational: row-batch chunking.
            for start in range(0, len(data), OPERATIONAL_ROWS_PER_CHUNK):
                batch   = data[start : start + OPERATIONAL_ROWS_PER_CHUNK]
                end_row = start + len(batch)
                prefix  = (
                    f"[Document: {file_name}]"
                    f" [Section: Rows {start + 1}–{end_row}]"
                )
                lines = ["Columns: " + " | ".join(h for h in headers if h)]
                for row in batch:
                    row_str = " | ".join(
                        f"{headers[i]}: {row[i]}"
                        for i in range(min(len(headers), len(row)))
                        if row[i].strip()
                    )
                    if row_str:
                        lines.append(row_str)
                body = "\n".join(lines)
                for ct in _split_long_text(body, prefix):
                    if _is_valid_chunk(ct):
                        chunks.append(Chunk(
                            chunk_id=str(uuid.uuid4()),
                            doc_id=doc_id,
                            file_name=file_name,
                            file_type="csv",
                            relative_path=file_path,
                            chunk_index=chunk_index,
                            chunk_text=ct,
                            section_header=f"Rows {start + 1}–{end_row}",
                        ))
                        chunk_index += 1

        print(f"  ✓ {file_name}: {len(chunks)} CSV chunks ({len(data)} rows)")
    except Exception as exc:
        print(f"  ✗ {file_name}: {exc}")

    return chunks


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".xlsm", ".docx", ".doc", ".csv"}


def parse_file(
    file_path: str,
    spark,
    vision_endpoint: Optional[str] = None,
) -> list[Chunk]:
    file_name = os.path.basename(file_path)
    ext       = Path(file_name).suffix.lower()
    doc_id    = make_doc_id(file_path)

    if ext == ".pdf":
        return parse_pdf(file_path, doc_id, file_name, spark, vision_endpoint=vision_endpoint)
    elif ext in {".xlsx", ".xls", ".xlsm"}:
        return parse_excel(file_path, doc_id, file_name)
    elif ext in {".docx", ".doc"}:
        return parse_word(file_path, doc_id, file_name)
    elif ext == ".csv":
        return parse_csv(file_path, doc_id, file_name)
    else:
        print(f"  — skipped unsupported type: {file_name}")
        return []


# ---------------------------------------------------------------------------
# Embedding helper
# ---------------------------------------------------------------------------

def get_embeddings_batch(texts: list[str], client, endpoint: str, batch_size: int = 20) -> list:
    """Generate embeddings in batches.

    Raises ValueError if any chunk exceeds MAX_CHUNK_CHARS — the upstream chunkers
    must enforce the limit; this function never truncates silently.
    """
    oversized = [(i, len(t)) for i, t in enumerate(texts) if len(t) > MAX_CHUNK_CHARS]
    if oversized:
        raise ValueError(
            f"{len(oversized)} chunk(s) exceed MAX_CHUNK_CHARS ({MAX_CHUNK_CHARS}): "
            f"{oversized[:5]}. Fix the upstream chunker."
        )
    embeddings = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        response = client.predict(endpoint=endpoint, inputs={"input": batch})
        embeddings.extend([item["embedding"] for item in response["data"]])
        if i % 100 == 0:
            print(f"  Embedded {i}/{len(texts)} chunks...")
    return embeddings


def _print_chunk_diagnostics(all_chunks: list) -> None:
    from collections import Counter
    if not all_chunks:
        print("\n=== Chunk diagnostics: no chunks ===")
        return
    sizes   = [c.char_count for c in all_chunks]
    by_type = Counter(c.file_type for c in all_chunks)
    by_file = Counter(c.file_name for c in all_chunks)
    print("\n=== Chunk diagnostics ===")
    print(f"Total chunks  : {len(all_chunks):,}")
    print(f"By file type  : {dict(by_type)}")
    print(f"Size — min    : {min(sizes):,}  max: {max(sizes):,}  avg: {sum(sizes)//len(sizes):,}")
    print(f"Oversized (>{MAX_CHUNK_CHARS:,}): {sum(1 for s in sizes if s > MAX_CHUNK_CHARS)}")
    print("Top 5 files by chunk count:")
    for fname, count in by_file.most_common(5):
        file_chunks = [c for c in all_chunks if c.file_name == fname]
        avg = sum(c.char_count for c in file_chunks) // count
        print(f"  {fname[:60]}: {count} chunks, avg {avg:,} chars")


# ---------------------------------------------------------------------------
# Vector search sync helper
# ---------------------------------------------------------------------------

def _wait_for_index_sync(
    spark,
    catalog: str,
    schema: str,
    index_suffix: str,
    table_embeddings: str,
    poll_interval: int = 30,
) -> None:
    """Trigger a Delta Sync on the vector index and block until it is confirmed
    current with the embeddings table.

    Strategy:
    1. Trigger sync.
    2. Obtain the DLT pipeline_id from the index spec.
    3. Poll the pipeline's latest update state until it reaches a terminal state
       (COMPLETED, FAILED, or CANCELED). This is the only reliable signal that
       the sync run has actually finished — idx.status.ready stays True throughout.
    4. After the pipeline finishes, compare indexed_row_count against the
       total row count in the embeddings table to confirm the index is current.
    5. Print a clear ✓ / ⚠ summary so the notebook output is unambiguous.
    """
    import time
    from databricks.sdk import WorkspaceClient

    index_name = f"{catalog}.{schema}.{index_suffix}"
    _TERMINAL_STATES = {"COMPLETED", "FAILED", "CANCELED"}

    try:
        w = WorkspaceClient()

        # Count total embeddings now in the source table.
        total_emb = spark.sql(f"SELECT COUNT(*) AS n FROM {table_embeddings}").collect()[0]["n"]

        # Trigger sync.
        w.vector_search_indexes.sync_index(index_name=index_name)
        print(f"\nVector search sync triggered → {index_name}")

        # Obtain pipeline_id from the index spec (needed to poll the DLT update).
        idx_info   = w.vector_search_indexes.get_index(index_name=index_name)
        pipeline_id = (
            idx_info.delta_sync_index_spec.pipeline_id
            if idx_info.delta_sync_index_spec
            else None
        )

        if not pipeline_id:
            print("  ⚠ Could not obtain pipeline_id — falling back to row-count polling.")
            pipeline_id = None

        print(f"  DLT pipeline : {pipeline_id or 'unknown'}")
        print(f"  Embeddings   : {total_emb:,} rows in source table")
        print(f"  Polling every {poll_interval}s ...\n")

        elapsed = 0
        while True:
            # Primary signal: DLT pipeline update state.
            if pipeline_id:
                try:
                    pipeline   = w.pipelines.get(pipeline_id=pipeline_id)
                    updates    = pipeline.latest_updates or []
                    state      = updates[0].state.value if updates else "PENDING"
                    state_str  = state
                except Exception:
                    state_str = "UNKNOWN"
            else:
                state_str = "UNKNOWN"

            # Secondary signal: indexed row count from the index status.
            idx_status   = w.vector_search_indexes.get_index(index_name=index_name)
            indexed_rows = idx_status.status.indexed_row_count or 0

            print(
                f"  [{elapsed:>4}s] pipeline={state_str:<12} "
                f"indexed={indexed_rows:,} / {total_emb:,}"
            )

            finished = state_str in _TERMINAL_STATES
            in_sync  = indexed_rows >= total_emb

            if finished and in_sync:
                print(f"\n✓ Index ready and current — {index_name}")
                print(f"  {indexed_rows:,} rows indexed  |  pipeline state: {state_str}")
                return

            if finished and not in_sync:
                # Pipeline completed but row count hasn't caught up yet — give it
                # one more cycle (replication can lag a few seconds).
                if elapsed > 0 and elapsed % (poll_interval * 3) == 0:
                    print(
                        f"  ⚠ Pipeline {state_str} but only {indexed_rows:,}/{total_emb:,} "
                        "rows indexed. Waiting for replication..."
                    )
                if state_str in ("FAILED", "CANCELED"):
                    print(f"  ✗ Sync pipeline ended with state: {state_str} — search may be stale.")
                    return

            time.sleep(poll_interval)
            elapsed += poll_interval

    except Exception as e:
        print(
            f"⚠ Could not sync vector index ({e}).\n"
            "  Run manually: w.vector_search_indexes.sync_index(index_name=...)\n"
            "  Do not proceed to semantic_search until the index is current."
        )


# ---------------------------------------------------------------------------
# Main workflow
# ---------------------------------------------------------------------------

def main():
    repo_root = find_repo_root()
    if repo_root not in sys.path:
        sys.path.insert(0, repo_root)
    print("repo_root:", repo_root)

    company_name       = get_param("sp_company_name")
    catalog            = get_param("catalog",  default="uc13")
    schema             = get_param("schema",   default="ingestion")
    embedding_endpoint = get_param("embedding_endpoint", default="databricks-bge-large-en")
    # Optional: a multimodal endpoint (e.g. databricks-meta-llama-3-2-11b-vision-instruct)
    # for extracting data from charts, org charts, and figures.  Leave empty to skip.
    _vision_raw    = get_param("vision_endpoint", default="")
    vision_endpoint: Optional[str] = _vision_raw.strip() or None
    if vision_endpoint:
        print(f"Vision endpoint : {vision_endpoint} (figure extraction enabled)")
    else:
        print("Vision endpoint : not configured (figure extraction skipped)")

    # parse_priority_tiers: "all" | "1" | "2" | "3" | "1,2" | "1,2,3" etc.
    parse_tiers_raw = get_param("parse_priority_tiers", default="all").strip().lower()
    if parse_tiers_raw == "all":
        tier_filter = ""
        tier_label  = "all tiers"
    else:
        tiers = [t.strip() for t in parse_tiers_raw.split(",") if t.strip().isdigit()]
        tier_filter = f"AND priority_tier IN ({', '.join(tiers)})"
        tier_label  = f"tier(s) {', '.join(tiers)}"

    volume_path      = f"/Volumes/{catalog}/{schema}/raw_files/{company_name}"
    table_relevance  = f"{catalog}.classification.doc_relevance"
    table_chunks     = f"{catalog}.{schema}.chunks"
    table_embeddings = f"{catalog}.{schema}.embeddings"

    from pyspark.sql import SparkSession as _SparkSession
    _spark = _SparkSession.getActiveSession()
    if _spark is None:
        raise RuntimeError("No active Spark session. This script must run on a Databricks cluster.")

    print(f"\n=== UC13 Phase 2b — Ingestion Parser ({company_name}) ===")
    print(f"Volume     : {volume_path}")
    print(f"Parsing    : {tier_label}")

    # --- Ensure output tables exist ---
    _spark.sql(f"CREATE SCHEMA IF NOT EXISTS {catalog}.{schema}")
    _spark.sql(f"""
        CREATE TABLE IF NOT EXISTS {table_chunks} (
            company_name   STRING,
            chunk_id       STRING,
            doc_id         STRING,
            file_name      STRING,
            file_type      STRING,
            relative_path  STRING,
            chunk_index    INT,
            chunk_text     STRING,
            section_header STRING,
            page_start     INT,
            page_end       INT,
            tab            STRING,
            source_type    STRING,
            char_count     INT,
            created_at     TIMESTAMP
        ) USING DELTA
    """)
    _spark.sql(f"""
        CREATE TABLE IF NOT EXISTS {table_embeddings} (
            company_name  STRING,
            chunk_id      STRING NOT NULL,
            doc_id        STRING,
            file_name     STRING,
            source_type   STRING,
            workstream    ARRAY<STRING>,
            priority_tier INT,
            embedding     ARRAY<FLOAT>,
            created_at    TIMESTAMP
        ) USING DELTA
        TBLPROPERTIES (
            'delta.enableChangeDataFeed'         = 'true',
            'delta.deletedFileRetentionDuration' = 'interval 30 days'
        )
    """)

    # --- Read files approved by classifier, lowest tier number first (1 = highest value) ---
    approved_rows = _spark.sql(f"""
        SELECT filename AS file_name, folder_path, workstream, priority_tier
        FROM {table_relevance}
        WHERE should_parse = true
          AND company_name = '{company_name}'
          {tier_filter}
        ORDER BY priority_tier ASC NULLS LAST
    """).collect()

    relevance_map = {
        r.file_name: {"workstream": list(r.workstream or []), "priority_tier": r.priority_tier}
        for r in approved_rows
    }

    file_paths = [
        os.path.join(volume_path, row.folder_path, row.file_name)
        if row.folder_path not in ("", ".", None)
        else os.path.join(volume_path, row.file_name)
        for row in approved_rows
    ]
    file_paths = [
        p for p in file_paths
        if os.path.exists(p) and Path(p).suffix.lower() in _ALLOWED_EXTENSIONS
    ]
    print(f"Files to parse: {len(file_paths)}")

    # --- Parse ---
    all_chunks: list[Chunk] = []
    for file_path in file_paths:
        chunks = parse_file(file_path, _spark, vision_endpoint=vision_endpoint)
        all_chunks.extend(chunks)

    _print_chunk_diagnostics(all_chunks)

    if not all_chunks:
        print("No chunks generated — exiting.")
        return

    # --- Save chunks ---
    from pyspark.sql import Row
    from pyspark.sql.types import (
        StructType, StructField, StringType, IntegerType, BooleanType,
        ArrayType, FloatType, TimestampType,
    )

    now = datetime.now(timezone.utc)

    chunk_schema = StructType([
        StructField("company_name",   StringType(),  False),
        StructField("chunk_id",       StringType(),  False),
        StructField("doc_id",         StringType(),  False),
        StructField("file_name",      StringType(),  False),
        StructField("file_type",      StringType(),  False),
        StructField("relative_path",  StringType(),  False),
        StructField("chunk_index",    IntegerType(), False),
        StructField("chunk_text",     StringType(),  False),
        StructField("section_header", StringType(),  True),
        StructField("page_start",     IntegerType(), True),
        StructField("page_end",       IntegerType(), True),
        StructField("tab",            StringType(),  True),
        StructField("source_type",    StringType(),  True),
        StructField("char_count",     IntegerType(), False),
        StructField("created_at",     TimestampType(), False),
    ])

    chunk_rows = [
        Row(
            company_name=company_name,
            chunk_id=c.chunk_id, doc_id=c.doc_id, file_name=c.file_name,
            file_type=c.file_type, relative_path=c.relative_path,
            chunk_index=int(c.chunk_index), chunk_text=c.chunk_text,
            section_header=c.section_header,
            page_start=int(c.page_start) if c.page_start is not None else None,
            page_end=int(c.page_end) if c.page_end is not None else None,
            tab=c.tab, source_type=c.source_type, char_count=int(c.char_count),
            created_at=now,
        )
        for c in all_chunks
    ]
    df_chunks = _spark.createDataFrame(chunk_rows, schema=chunk_schema)
    # Replace this company's chunks so re-runs are idempotent.
    try:
        _spark.sql(f"DELETE FROM {table_chunks} WHERE company_name = '{company_name}'")
    except Exception:
        pass
    df_chunks.write.mode("append").option("mergeSchema", "true").saveAsTable(table_chunks)
    print(f"✓ Saved {df_chunks.count()} chunks → {table_chunks}")

    # --- Generate and save embeddings ---
    import mlflow.deployments

    client = mlflow.deployments.get_deploy_client("databricks")
    texts  = [c.chunk_text for c in all_chunks]
    print(f"\nGenerating embeddings for {len(texts)} chunks...")
    embeddings = get_embeddings_batch(texts, client, embedding_endpoint)
    print(f"Generated {len(embeddings)} embeddings")

    emb_schema = StructType([
        StructField("company_name",  StringType(),           False),
        StructField("chunk_id",      StringType(),           False),
        StructField("doc_id",        StringType(),           False),
        StructField("file_name",     StringType(),           False),
        StructField("source_type",   StringType(),           True),
        StructField("workstream",    ArrayType(StringType()), True),
        StructField("priority_tier", IntegerType(),          True),
        StructField("embedding",     ArrayType(FloatType()), False),
        StructField("created_at",    TimestampType(),        False),
    ])

    emb_rows = [
        Row(
            company_name=company_name,
            chunk_id=all_chunks[i].chunk_id,
            doc_id=all_chunks[i].doc_id,
            file_name=all_chunks[i].file_name,
            source_type=all_chunks[i].source_type,
            workstream=relevance_map.get(all_chunks[i].file_name, {}).get("workstream"),
            priority_tier=relevance_map.get(all_chunks[i].file_name, {}).get("priority_tier"),
            embedding=[float(x) for x in embeddings[i]],
            created_at=now,
        )
        for i in range(len(all_chunks))
    ]
    df_emb = _spark.createDataFrame(emb_rows, schema=emb_schema)
    # Replace this company's embeddings so re-runs are idempotent.
    try:
        _spark.sql(f"DELETE FROM {table_embeddings} WHERE company_name = '{company_name}'")
    except Exception:
        pass
    df_emb.write.mode("append").option("mergeSchema", "true").saveAsTable(table_embeddings)
    print(f"✓ Saved {df_emb.count()} embeddings → {table_embeddings}")

    # --- Trigger vector search index sync and wait for completion ---
    _wait_for_index_sync(
        spark=_spark,
        catalog=catalog,
        schema=schema,
        index_suffix="embeddings_index",
        table_embeddings=table_embeddings,
    )


if __name__ == "__main__":
    main()
