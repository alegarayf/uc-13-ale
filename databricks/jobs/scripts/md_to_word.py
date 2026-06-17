"""
md_to_word.py — Convert PE diligence markdown assessments to styled Word documents.

Handles both Financial Trends and Business Model assessment report formats.
python-docx is already installed in the cluster environment (Cell 0 pip install).

Usage from test_pipeline.ipynb:
    import sys
    sys.path.insert(0, repo_root)
    from jobs.scripts.md_to_word import convert_md_to_word

    convert_md_to_word(
        md_path  = "/Volumes/uc13/analysis/reports/Elder_Care/financial_trends_report.md",
        out_path = "/Volumes/uc13/analysis/reports/Elder_Care/financial_trends_report.docx",
    )

Or as a standalone script:
    python md_to_word.py input.md output.docx
"""

import re
import shutil
import sys
import tempfile
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Color palette ──────────────────────────────────────────────────────────────
_NAVY       = "1F3864"   # dark navy    — title, H2, table headers
_BLUE       = "2E75B6"   # medium blue  — H3, sub-headers
_ROW_ALT    = "EEF3FA"   # light blue   — alternating data rows
_RED_BG     = "FFE5E5"   # light red    — red flag rows/cells
_RED_TXT    = "C00000"   # dark red     — red flag text
_YLW_BG     = "FFF2CC"   # light amber  — yellow flag rows, ⚠️ warnings
_YLW_TXT    = "7F6000"   # dark amber   — yellow flag text
_BOX_BG     = "F2F4F7"   # light grey   — analyst take, executive summary callouts
_META_TXT   = "595959"   # dark grey    — metadata, italic body
_BORDER     = "BFD0E8"   # blue-grey    — table borders

# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _hex_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_bg(cell, color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.upper())
    tcPr.append(shd)


def _cell_valign(cell, align: str = "center"):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), align)
    tcPr.append(el)


def _para_shading(para, color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.upper())
    pPr.append(shd)


def _clear_table_borders(table):
    """Remove all grid borders from a table (for accounting-style P&L format)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _cell_bottom_border(cell, val: str = "single", size: str = "8",
                        color: str = "000000"):
    """Add a bottom border to a single cell (for subtotal rows in P&L tables)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for old in tcBorders.findall(qn("w:bottom")):
        tcBorders.remove(old)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color.upper())
    tcBorders.append(bottom)


def _table_borders(table, color: str = _BORDER):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _row_height(row, pt: float):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(pt * 20)))
    trPr.append(h)


def _para_bottom_border(para, color: str = _NAVY, size: str = "8"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _start_portrait_section(doc):
    """Paragraph that ENDS a portrait section (content before a wide table)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "12240")   # 8.5"
    pgSz.set(qn("w:h"), "15840")   # 11"
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "1080"); pgMar.set(qn("w:bottom"), "1080")
    pgMar.set(qn("w:left"), "1080"); pgMar.set(qn("w:right"), "1080")
    sectPr.append(pgSz); sectPr.append(pgMar)
    pPr.append(sectPr)


def _end_landscape_section(doc):
    """Paragraph that ENDS the landscape section (wide table is contained here)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "15840")   # 11"
    pgSz.set(qn("w:h"), "12240")   # 8.5"
    pgSz.set(qn("w:orient"), "landscape")
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "720"); pgMar.set(qn("w:bottom"), "720")
    pgMar.set(qn("w:left"), "720"); pgMar.set(qn("w:right"), "720")
    sectPr.append(pgSz); sectPr.append(pgMar)
    pPr.append(sectPr)


# ── Inline markdown parser ─────────────────────────────────────────────────────

def _inline_runs(para, text: str, size: float = 9.5,
                 color: str = "000000", bold: bool = False, italic: bool = False):
    """Parse **bold**, *italic*, _italic_ inline markers and add styled runs."""
    pattern = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|_[^_]+?_)")
    for part in pattern.split(text):
        if not part:
            continue
        run = para.add_run()
        run.font.size = Pt(size)
        run.font.color.rgb = _hex_rgb(color)
        run.bold = bold
        run.italic = italic
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run.text = part[2:-2]
            run.bold = True
        elif (part.startswith("*") and part.endswith("*")) or \
             (part.startswith("_") and part.endswith("_")) and len(part) > 2:
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part


# ── Markdown block parser ──────────────────────────────────────────────────────

def _parse_md(text: str) -> List[dict]:
    blocks = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings (order matters — check h3/h2 before h1)
        if line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
            i += 1; continue
        if line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
            i += 1; continue
        if line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
            i += 1; continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            blocks.append({"type": "divider"})
            i += 1; continue

        # Blockquote (executive summary, inline notes)
        if line.startswith("> "):
            blocks.append({"type": "quote", "text": line[2:].strip()})
            i += 1; continue

        # Pipe table — accumulate consecutive rows
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                raw = lines[i]
                cells = [c.strip() for c in raw.strip("|").split("|")]
                # Skip separator rows (cells are all dashes/colons/spaces)
                if not all(re.match(r"^[-:| ]*$", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # Bullet list — collect consecutive items
        if re.match(r"^[-*] ", line):
            items = []
            while i < len(lines) and re.match(r"^[-*] ", lines[i]):
                items.append(lines[i][2:].strip())
                i += 1
            if items:
                blocks.append({"type": "bullets", "items": items})
            continue

        # ⚠️ warning line
        if line.startswith("⚠️"):
            blocks.append({"type": "warning", "text": line.strip()})
            i += 1; continue

        # Analyst take paragraph
        if re.match(r"^\*\*Analyst take:\*\*", line):
            blocks.append({"type": "analyst_take", "text": line.strip()})
            i += 1; continue

        # Generated/metadata line (bold + pipe separators)
        if re.match(r"^\*\*Generated", line) or ("**Generated**" in line):
            blocks.append({"type": "meta", "text": line.strip()})
            i += 1; continue

        # Bold paragraph (entity structure, management depth, etc.)
        if line.startswith("**") and (":**" in line or line.endswith("**")):
            blocks.append({"type": "bold_para", "text": line.strip()})
            i += 1; continue

        # Blank / empty
        if not line.strip():
            blocks.append({"type": "blank"})
            i += 1; continue

        # Regular paragraph
        blocks.append({"type": "para", "text": line.strip()})
        i += 1
    return blocks


# ── Document setup ─────────────────────────────────────────────────────────────

def _new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = int(Inches(8.5))
    sec.page_height   = int(Inches(11))
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    # Base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    return doc


# ── Block renderers ────────────────────────────────────────────────────────────

def _r_h1(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _hex_rgb(_NAVY)


def _r_h2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _hex_rgb(_NAVY)
    _para_bottom_border(p, color=_NAVY, size="8")


def _r_h3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _hex_rgb(_BLUE)


def _r_meta(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(8)
    cleaned = re.sub(r"\*\*", "", text)
    run = p.add_run(cleaned)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _hex_rgb(_META_TXT)


def _r_quote(doc, text: str):
    """Executive summary / blockquote — grey shaded callout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    _para_shading(p, _BOX_BG)
    _inline_runs(p, text, size=9.5, italic=True, color=_META_TXT)


def _r_analyst_take(doc, text: str):
    """'Analyst take:' callout in grey box."""
    body = re.sub(r"^\*\*Analyst take:\*\*\s*", "", text).strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    _para_shading(p, _BOX_BG)
    lbl = p.add_run("Analyst take: ")
    lbl.bold = True; lbl.italic = True
    lbl.font.size = Pt(9.5)
    lbl.font.color.rgb = _hex_rgb(_NAVY)
    _inline_runs(p, body, size=9.5, italic=True, color=_META_TXT)


def _r_warning(doc, text: str):
    """⚠️ EBITDA warning line — amber shaded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.1)
    _para_shading(p, _YLW_BG)
    _inline_runs(p, text, size=9.5, color=_YLW_TXT)


def _r_bold_para(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    _inline_runs(p, text, size=9.5)


def _r_para(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _inline_runs(p, text, size=9.5)


def _detect_severity(text: str) -> Optional[str]:
    if "🔴" in text:
        return "red"
    if "🟡" in text:
        return "yellow"
    if "🟢" in text:
        return "green"
    return None


def _r_bullets(doc, items: List[str]):
    """Flag-aware bullet list. Flag items (🔴/🟡) render as a compact table."""
    is_flags = all(_detect_severity(it) for it in items) and items
    if is_flags:
        _r_flag_table(doc, items)
        return
    for item in items:
        sev = _detect_severity(item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        if sev == "red":
            _para_shading(p, _RED_BG)
        elif sev == "yellow":
            _para_shading(p, _YLW_BG)
        txt_color = _RED_TXT if sev == "red" else (_YLW_TXT if sev == "yellow" else "000000")
        bullet_run = p.add_run("• ")
        bullet_run.font.size = Pt(9.5)
        bullet_run.font.color.rgb = _hex_rgb(txt_color)
        _inline_runs(p, item, size=9.5, color=txt_color)


def _r_flag_table(doc, items: List[str]):
    """Render investment flag bullets as a 3-col table: Severity | Metric & Value | Note."""
    rows_data = []
    for item in items:
        sev = _detect_severity(item)
        sev_label = ("🔴 Red" if sev == "red" else
                     "🟡 Yellow" if sev == "yellow" else
                     "🟢 Green")
        body = re.sub(r"^[🔴🟡🟢]\s*", "", item).strip()
        if " — " in body:
            metric_val, note = body.split(" — ", 1)
        else:
            metric_val, note = body, ""
        rows_data.append((sev, sev_label, metric_val.strip(), note.strip()))

    n = len(rows_data) + 1
    table = doc.add_table(rows=n, cols=3)
    table.style = "Table Grid"
    _table_borders(table)

    # Header
    hdrs = ["Severity", "Metric / Value", "Note"]
    for col_idx, hdr in enumerate(hdrs):
        cell = table.cell(0, col_idx)
        _cell_bg(cell, _NAVY)
        _cell_valign(cell, "center")
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = _hex_rgb("FFFFFF")

    # Data rows
    for row_idx, (sev, sev_label, metric_val, note) in enumerate(rows_data):
        r = row_idx + 1
        bg = _RED_BG if sev == "red" else (_YLW_BG if sev == "yellow" else "FFFFFF")
        tc = _YLW_TXT if sev == "yellow" else (_RED_TXT if sev == "red" else _NAVY)

        cells_data = [sev_label, metric_val, note]
        for col_idx, txt in enumerate(cells_data):
            cell = table.cell(r, col_idx)
            _cell_bg(cell, bg)
            _cell_valign(cell, "center")
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.05)
            _inline_runs(p, txt, size=9.0, color=tc,
                         bold=(col_idx == 0))

    # Column widths: Severity=1.0, Metric=2.8, Note=3.2
    for row in table.rows:
        for col_idx, w in enumerate([1.0, 2.8, 3.2]):
            row.cells[col_idx].width = int(Inches(w))
    _row_height(table.rows[0], 14)
    doc.add_paragraph()


def _is_pl_table(rows: List[List[str]]) -> bool:
    """True when the table looks like a period-column P&L (first header = 'Line Item')."""
    if not rows or len(rows[0]) < 2:
        return False
    first = (rows[0][0] or "").strip().lower()
    if first != "line item":
        return False
    period_re = re.compile(r"\d{4}|ttm|ltm", re.I)
    return any(period_re.search(h) for h in rows[0][1:])


def _r_pl_table(doc: Document, rows: List[List[str]]):
    """Render a P&L table in clean accounting format — no grid borders,
    right-aligned numbers, bold subtotals with thin underlines.

    Mirrors the Amazon / Walmart income-statement style:
    - White background, no cell shading
    - Thin bottom border on every bold (subtotal) row
    - Double bottom border on the very last bold row
    - ↳ rows indented in the label column
    - Numbers right-aligned; labels left-aligned
    - Header row: bold, light grey background, thin bottom border
    """
    if not rows:
        return

    n_cols = len(rows[0])
    n_data_cols = max(n_cols - 1, 1)

    # Column widths — label gets more room, data cols share the rest
    avail_w = 7.0
    label_w = min(2.8, avail_w * 0.40)
    data_col_w = round((avail_w - label_w) / n_data_cols, 3)
    data_col_w = max(0.65, min(data_col_w, 1.3))

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Normal"
    table.allow_autofit = False
    _clear_table_borders(table)

    # Identify bold (subtotal) rows — raw cell text starts with **
    def _cell_is_bold(text: str) -> bool:
        t = (text or "").strip()
        return t.startswith("**") and t.endswith("**") and len(t) > 4

    bold_rows = [i for i, r in enumerate(rows) if r and _cell_is_bold(r[0])]
    last_bold = bold_rows[-1] if bold_rows else None

    for row_idx, row_cells in enumerate(rows):
        is_header = (row_idx == 0)
        label_raw = row_cells[0] if row_cells else ""
        is_bold   = _cell_is_bold(label_raw) and not is_header
        is_last_bold = (row_idx == last_bold)
        is_indented  = (label_raw or "").strip().startswith("↳")

        tr = table.rows[row_idx]
        _row_height(tr, 15 if is_header else 13)

        for col_idx in range(n_cols):
            cell_text = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell = table.cell(row_idx, col_idx)
            _cell_valign(cell, "center")

            # Column width
            cell.width = int(Inches(label_w if col_idx == 0 else data_col_w))

            # Header: light grey tint + thin bottom line
            if is_header:
                _cell_bg(cell, "F2F4F7")
                _cell_bottom_border(cell, "single", size="8", color=_NAVY)

            # Bold (subtotal) row: thin bottom underline
            elif is_bold:
                border_type = "double" if is_last_bold else "single"
                _cell_bottom_border(cell, border_type, size="6", color="404040")

            # ── Text ────────────────────────────────────────────────────────
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

            # Label column — left-aligned with optional indent
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if is_indented:
                    p.paragraph_format.left_indent = Inches(0.22)
                else:
                    p.paragraph_format.left_indent = Inches(0.04)
                color = _NAVY if (is_header or is_bold) else "1A1A1A"
                _inline_runs(p, cell_text, size=9.5, color=color,
                             bold=(is_header or is_bold))

            # Data columns — right-aligned
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_indent = Inches(0.06)
                txt = (cell_text or "").strip()
                # Strip bold markers from data cells (they're bolded via run style)
                if txt.startswith("**") and txt.endswith("**"):
                    txt = txt[2:-2]
                color = _NAVY if (is_header or is_bold) else "1A1A1A"
                run = p.add_run(txt)
                run.font.size = Pt(9.5)
                run.bold = is_header or is_bold
                run.font.color.rgb = _hex_rgb(color)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(4)


def _r_table(doc, rows: List[List[str]]):
    """Render a markdown pipe table as a styled Word table.

    P&L summary tables (first header cell = 'Line Item') are routed through
    _r_pl_table() for clean accounting-style formatting (no grid borders,
    right-aligned numbers, bold subtotals with underlines).
    All other tables use the colored grid style.
    """
    if not rows:
        return
    if _is_pl_table(rows):
        _r_pl_table(doc, rows)
        return
    n_cols = len(rows[0])
    n_rows = len(rows)
    is_wide = n_cols > 8

    if is_wide:
        _start_portrait_section(doc)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.allow_autofit = False
    _table_borders(table)

    # Available width
    avail = 9.5 if is_wide else 7.0   # inches (landscape 11"-1.5"margin vs portrait)
    if n_cols > 1:
        label_w = 1.8 if is_wide else min(2.4, avail * 0.32)
        col_w   = (avail - label_w) / (n_cols - 1)
        widths  = [label_w] + [col_w] * (n_cols - 1)
    else:
        widths = [avail]

    font_size = 7.5 if is_wide else 9.0

    for row_idx, row_cells in enumerate(rows):
        is_header = (row_idx == 0)
        tr = table.rows[row_idx]
        _row_height(tr, 14 if is_header else 13)

        for col_idx in range(n_cols):
            cell_text = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell = table.cell(row_idx, col_idx)
            _cell_valign(cell, "center")

            # Background
            sev = _detect_severity(cell_text) if not is_header else None
            if is_header:
                _cell_bg(cell, _NAVY)
            elif sev == "red":
                _cell_bg(cell, _RED_BG)
            elif sev == "yellow":
                _cell_bg(cell, _YLW_BG)
            elif row_idx % 2 == 0:
                _cell_bg(cell, _ROW_ALT)

            # Column width
            if col_idx < len(widths):
                cell.width = int(Inches(widths[col_idx]))

            # Text rendering
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

            # ↳ prefix in the label column = indented sub-row (segment, sub-metric).
            txt = cell_text
            indent_extra = Inches(0.05)
            if col_idx == 0 and txt.startswith("↳"):
                indent_extra = Inches(0.20)
            p.paragraph_format.left_indent = indent_extra

            txt_color = ("FFFFFF" if is_header else
                         _RED_TXT if sev == "red" else
                         _YLW_TXT if sev == "yellow" else "000000")
            _inline_runs(p, txt, size=font_size, color=txt_color, bold=is_header)

    if is_wide:
        _end_landscape_section(doc)

    # Small breathing room
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after  = Pt(2)


def _r_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _para_bottom_border(p, color=_NAVY, size="4")


# ── Dispatcher ─────────────────────────────────────────────────────────────────

def _render(doc: Document, block: dict):
    t = block["type"]
    if t == "h1":
        _r_h1(doc, block["text"])
    elif t == "h2":
        _r_h2(doc, block["text"])
    elif t == "h3":
        _r_h3(doc, block["text"])
    elif t == "meta":
        _r_meta(doc, block["text"])
    elif t == "quote":
        _r_quote(doc, block["text"])
    elif t == "analyst_take":
        _r_analyst_take(doc, block["text"])
    elif t == "warning":
        _r_warning(doc, block["text"])
    elif t == "bullets":
        _r_bullets(doc, block["items"])
    elif t == "table":
        _r_table(doc, block["rows"])
    elif t == "bold_para":
        _r_bold_para(doc, block["text"])
    elif t == "para":
        _r_para(doc, block["text"])
    elif t == "divider":
        _r_divider(doc)
    # blank: skip


# ── Public API ─────────────────────────────────────────────────────────────────

def convert_md_to_word(md_path: str, out_path: str) -> str:
    """Convert a PE diligence markdown report to a styled Word (.docx) file.

    Args:
        md_path:  Path to the source .md file.
        out_path: Destination path for the .docx file (parent dir created if needed).

    Returns:
        out_path — for display/chaining in notebook cells.
    """
    text   = Path(md_path).read_text(encoding="utf-8")
    blocks = _parse_md(text)
    doc    = _new_doc()

    prev_type = None
    for block in blocks:
        btype = block["type"]
        # Suppress consecutive or leading blanks after structural elements
        if btype == "blank":
            if prev_type in (None, "blank", "divider", "h1", "h2", "h3"):
                continue
        _render(doc, block)
        prev_type = btype

    # UC Volumes don't support the random-write seek() calls that python-docx's
    # underlying zipfile needs to finalize the .docx ZIP structure. Save to a
    # local /tmp file first, then copy to the destination (works on any path).
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as _tmp:
        _tmp_path = _tmp.name
    try:
        doc.save(_tmp_path)
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(_tmp_path, str(out))
    finally:
        Path(_tmp_path).unlink(missing_ok=True)

    print(f"✓ Word document saved → {out_path}")
    return out_path


# ── CLI entry point ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md_to_word.py <input.md> <output.docx>")
        sys.exit(1)
    convert_md_to_word(sys.argv[1], sys.argv[2])
